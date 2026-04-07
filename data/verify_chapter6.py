"""Verify all corrections in chapter6_results_v3.docx"""
from docx import Document
doc = Document('chapter6_results_v3.docx')

checks = [
    ('72 genuine click events', 'Traceability total fix'),
    ('early in his session', 'Shanyl pronoun fix'),
    ('correlates most closely', 'Causal overclaim fix'),
    ('All identified disparities', 'Disparities framing fix'),
    ('confirmed instrumentation bug that misattributed', 'Footnote component strip'),
    ('reply target reference was captured', 'parentMessageId strip (para 9)'),
    ('asynchronous research pipeline', 'Controller names strip'),
    ('structured report generation pathway', 'generateReport strip'),
    ('34 context promotion actions', 'Event name strip (para 46)'),
    ('21 context promotion actions in solo', 'Event name strip (para 75 promotes)'),
    ('panel-to-chat navigation click dominance', 'Event name strip (para 75 jump)'),
    ('four AI toggle changes', 'Event name strip (para 75 toggle)'),
    ('insight status component logged all actions under a default', 'Component strip (para 77)'),
    ('20 insight status change events', 'Event name strip (para 77)'),
    ('reply target reference was not forwarded', 'parentMessageId strip (para 81)'),
    ('direct generation events for S1 and S4', 'Pipeline strip (para 81)'),
]

all_text = ' '.join([p.text for p in doc.paragraphs])
print('=== VERIFICATION ===')
for needle, label in checks:
    found = needle in all_text
    print(f'  {"OK" if found else "FAIL"} {label}')

gone = [
    ('InsightActions component', 'InsightActions removed'),
    ('ResearchJobController', 'ResearchJobController removed'),
    ('AIInsightController synchronous', 'AIInsightController removed'),
    ('generateReport with the research', 'generateReport removed'),
    ('draft_context_promoted', 'draft_context_promoted removed'),
    ('jump_to_chat_marker event', 'jump_to_chat_marker removed'),
    ('team_ai_toggle_changed', 'team_ai_toggle_changed removed'),
    ('insight_status_changed events recorded', 'insight_status_changed removed'),
    ('parentMessageId', 'parentMessageId removed'),
    ('55 genuine click', '55 count removed'),
    ('determining factor was not', 'Causal overclaim removed'),
    ('in her session', 'Wrong pronoun removed'),
    ('participant-generated adaptation confirmed by 21', 'False conflation removed'),
]

print()
print('=== REMOVALS ===')
for needle, label in gone:
    absent = needle not in all_text
    print(f'  {"OK" if absent else "FAIL"} {label}')

table = doc.tables[0]
s1 = table.cell(1, 4).text.strip()
total = table.cell(5, 4).text.strip()
print()
print('=== TABLE ===')
print(f'  S1 traceability: "{s1}" {"OK" if s1 == "22*" else "FAIL"}')
print(f'  Total traceability: "{total}" {"OK" if total == "72" else "FAIL"}')
