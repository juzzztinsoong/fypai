"""
Rewrite script for chapter6_results_v2.docx
Applies factual corrections, framing fixes, and strips implementation-level detail.
"""
from docx import Document

doc = Document('chapter6_results_v2.docx')


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """Replace old_text with new_text in paragraph, preserving run-level formatting."""
    full_text = ''.join([run.text for run in paragraph.runs])
    if old_text not in full_text:
        return False

    # Build character-to-run mapping
    char_map = []  # (run_index, char_index_within_run)
    for ri, run in enumerate(paragraph.runs):
        for ci in range(len(run.text)):
            char_map.append((ri, ci))

    start = full_text.index(old_text)
    end = start + len(old_text)

    if start_run_idx := char_map[start][0] == char_map[end - 1][0]:
        # Single-run case
        ri = char_map[start][0]
        sc = char_map[start][1]
        ec = char_map[end - 1][1]
        run = paragraph.runs[ri]
        run.text = run.text[:sc] + new_text + run.text[ec + 1:]
    else:
        # Cross-run case
        start_run, start_char = char_map[start]
        end_run, end_char = char_map[end - 1]
        # Put replacement text at end of start_run's prefix
        paragraph.runs[start_run].text = paragraph.runs[start_run].text[:start_char] + new_text
        # Clear middle runs
        for ri in range(start_run + 1, end_run):
            paragraph.runs[ri].text = ''
        # Keep end_run's suffix
        paragraph.runs[end_run].text = paragraph.runs[end_run].text[end_char + 1:]

    return True


def replace_in_doc(old_text, new_text, label=""):
    """Try to replace in all paragraphs. Returns count of replacements."""
    count = 0
    for para in doc.paragraphs:
        if old_text in ''.join([r.text for r in para.runs]):
            if replace_text_in_paragraph(para, old_text, new_text):
                count += 1
    status = "OK" if count else "MISS"
    short = label or old_text[:60]
    print(f"  {status} ({count}): {short}")
    return count


print("=== FACTUAL CORRECTIONS ===")

# 1. Para 27: Wrong traceability total (55 should be 72)
replace_in_doc(
    "55 genuine click events across nine participants",
    "72 genuine click events across nine participants",
    "Fix traceability total: 55 -> 72"
)

# 2. Para 23: Shanyl is male, not female
replace_in_doc(
    "early in her session, preventing the contrast from forming",
    "early in his session, preventing the contrast from forming",
    "Fix Shanyl pronoun: her -> his"
)

# 3. Para 19: Conflates reply workaround with draft promotes (different mechanisms)
replace_in_doc(
    " \u2014 a participant-generated adaptation confirmed by 21 draft_context_promoted events in the solo phase",
    "",
    "Remove false conflation of reply workaround with draft promotes"
)

print("\n=== FRAMING FIXES ===")

# 4. Para 23: Overclaims causality with n=7
replace_in_doc(
    "The determining factor was not general AI literacy but depth of Research mode engagement",
    "The pattern correlates most closely with depth of Research mode engagement rather than general AI literacy",
    "Fix causal overclaim -> correlation"
)

# 5. Para 75: Overly absolute claim
replace_in_doc(
    "No unresolved disparities between the log data and cleaned observational notes were identified.",
    "All identified disparities between the log data and cleaned observational notes were resolved.",
    "Soften 'no unresolved disparities'"
)

print("\n=== STRIP IMPLEMENTATION DETAIL ===")

# 6. Para 7 (footnote): Component name
replace_in_doc(
    "a confirmed userId bug in the InsightActions component",
    "a confirmed instrumentation bug that misattributed the actor identifier",
    "Strip InsightActions component name from footnote"
)

# 7. Para 9: parentMessageId
replace_in_doc(
    "where parentMessageId was extracted but not passed to the LLM",
    "where the reply target reference was captured but not forwarded to the language model",
    "Strip parentMessageId from para 9"
)

# 8. Para 9: Controller/pipeline names
replace_in_doc(
    "S1 used the ResearchJobController async pathway for research-type insights; S2 through S4 used the AIInsightController synchronous pathway with a research-analyst archetype modifier",
    "S1 used an asynchronous research pipeline with a deeper research-specific prompt; S2 through S4 used a synchronous generation pathway with a research-analyst framing",
    "Strip controller names from para 9"
)

# 9. Para 20: generateReport function name
replace_in_doc(
    "routes to generateReport with the research-analyst archetype modifier, operating",
    "routes to a structured report generation pathway, operating",
    "Strip generateReport from para 20"
)

# 10. Para 46: Event name
replace_in_doc(
    "34 draft_context_promoted events across all phases",
    "34 context promotion actions across all phases",
    "Strip event name from para 46"
)

# 11. Para 75: Event names (multiple)
replace_in_doc(
    "21 draft_context_promoted events in solo",
    "21 context promotion actions in solo",
    "Strip event name from para 75 (promotes)"
)
replace_in_doc(
    "jump_to_chat_marker event direction dominance",
    "panel-to-chat navigation click dominance",
    "Strip event name from para 75 (jump_to_chat)"
)
replace_in_doc(
    "four team_ai_toggle_changed events in 18 seconds",
    "four AI toggle changes in 18 seconds",
    "Strip event name from para 75 (toggle)"
)

# 12. Para 77: Component name and event name
replace_in_doc(
    "the InsightActions component defaulted userId to user1 rather than the session participant\u2019s actual ID",
    "the insight status component logged all actions under a default identifier rather than the session participant\u2019s actual ID",
    "Strip InsightActions from para 77 (smart quote)"
)
# Try with straight quote too
replace_in_doc(
    "the InsightActions component defaulted userId to user1 rather than the session participant's actual ID",
    "the insight status component logged all actions under a default identifier rather than the session participant's actual ID",
    "Strip InsightActions from para 77 (straight quote)"
)
replace_in_doc(
    "20 insight_status_changed events recorded",
    "20 insight status change events recorded",
    "Strip event name from para 77"
)

# 13. Para 81: Implementation details
replace_in_doc(
    "parentMessageId was not passed to the LLM",
    "the reply target reference was not forwarded to the language model",
    "Strip parentMessageId from para 81"
)
replace_in_doc(
    "timeline events for S1 and S4, export fallback for S2 and S3",
    "direct generation events for S1 and S4, export-based reconstruction for S2 and S3",
    "Strip pipeline names from para 81"
)

print("\n=== TABLE FIXES ===")

# Fix Table 5
table = doc.tables[0]

# Row 1 (S1): column 4 — "22 (15 genuine*)" -> "22*"
cell = table.cell(1, 4)
cell_text = cell.paragraphs[0].text if cell.paragraphs else ""
if "15 genuine" in cell_text or "(15" in cell_text:
    for run in cell.paragraphs[0].runs:
        run.text = ""
    cell.paragraphs[0].runs[0].text = "22*"
    print("  OK: Fixed S1 traceability cell -> 22*")
else:
    print(f"  MISS: S1 traceability cell (found: '{cell_text}')")

# Row 5 (Total): column 4 — "55+" -> "72"
cell = table.cell(5, 4)
cell_text = cell.paragraphs[0].text if cell.paragraphs else ""
if "55" in cell_text:
    for run in cell.paragraphs[0].runs:
        run.text = ""
    cell.paragraphs[0].runs[0].text = "72"
    print("  OK: Fixed Total traceability cell -> 72")
else:
    print(f"  MISS: Total traceability cell (found: '{cell_text}')")

# Save
output_path = 'chapter6_results_v3.docx'
doc.save(output_path)
print(f"\nSaved as {output_path}")
